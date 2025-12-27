from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Inches, Cm, Mm

TEX_INLINE_PATTERN = re.compile(r"\$(?P<body>[^$]+)\$")
TEX_BLOCK_INLINE_PATTERN = re.compile(r"\$\$(?P<body>[\s\S]+?)\$\$")
TEX_TEXT_COMMAND_PATTERN = re.compile(r"\\text\{([^}]*)\}")
TEX_COMMAND_PATTERN = re.compile(r"\\[A-Za-z]+")
TEX_FRACTION_PATTERN = re.compile(r"\\frac\{([^{}]+)\}\{([^{}]+)\}")
TEX_SUB_SUP_PATTERN = re.compile(r"([A-Za-z]+)\s*[_^]\s*\{?(\d+)\}?")

TABLE_RULE = re.compile(r"^:?-{3,}:?$")
IMG_HTML_PATTERN = re.compile(r"<img[^>]*src=\"([^\"]+)\"[^>]*>")
IMG_MD_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")
WIDTH_PATTERN = re.compile(r"width\s*=\s*\"?([0-9]+(?:\.[0-9]+)?)(px|cm|mm)?\"?")
PAGE_HEADING_PATTERN = re.compile(r"^#\s+Page\s+(?P<page>\d+)\s*$")


def read_markdown(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    return text.replace("\r\n", "\n").splitlines()


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.count("|") >= 2


def split_table_line(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def looks_like_divider(row: list[str]) -> bool:
    return row and all(TABLE_RULE.match(cell.replace(" ", "")) for cell in row)


def is_placeholder_row(row: list[str]) -> bool:
    for cell in row:
        stripped = cell.replace("<br>", "").strip()
        if stripped and stripped not in {"-", "–", "—"}:
            return False
    return True


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = [split_table_line(line) for line in table_lines if line.strip()]
    if not rows:
        return

    cleaned = []
    for row in rows:
        if looks_like_divider(row):
            continue
        if is_placeholder_row(row):
            continue
        cleaned.append(row)

    if not cleaned:
        cleaned = rows

    cols = max(len(row) for row in cleaned)
    table = document.add_table(rows=len(cleaned), cols=cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(cleaned):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            text = value.replace("<br>", "\n").strip()
            if text == "-":
                text = ""
            table.rows[r_idx].cells[c_idx].text = text


def flush_paragraph(document: Document, buffer: list[str], base_dir: Path | None = None) -> None:
    if not buffer:
        return
    text = " ".join(buffer).strip()
    buffer.clear()
    if not text:
        return
    paragraph = document.add_paragraph()
    render_text_content(paragraph, text)


def strip_tex_math_delimiters(text: str) -> str:
    """docx 出力向けに、TeX デリミタだけ除去して中身をそのまま残す。"""

    def strip_inline(match: re.Match[str]) -> str:
        return match.group("body")

    stripped = text.strip()
    if stripped == "$$":
        return ""

    text = text.replace("\\[", "").replace("\\]", "")
    text = text.replace("\\(", "").replace("\\)", "")
    text = TEX_BLOCK_INLINE_PATTERN.sub(lambda m: m.group("body").strip(), text)
    prev = None
    while prev != text:
        prev = text
        text = TEX_INLINE_PATTERN.sub(strip_inline, text)

    text = TEX_TEXT_COMMAND_PATTERN.sub(lambda m: m.group(1), text)
    text = TEX_FRACTION_PATTERN.sub(lambda m: f"({m.group(1)})/({m.group(2)})", text)
    text = TEX_SUB_SUP_PATTERN.sub(lambda m: f"{m.group(1)}_{m.group(2)}", text)
    text = text.replace("{", "").replace("}", "")
    text = TEX_COMMAND_PATTERN.sub("", text)
    return text


def render_text_content(paragraph, text: str) -> None:
    working = text.replace("<br>", "\n")
    working = strip_tex_math_delimiters(working)
    if not working:
        return
    paragraph.add_run(working)


def to_width(value: str | None):
    if not value:
        return None
    match = WIDTH_PATTERN.search(value)
    if not match:
        return None
    amount = float(match.group(1))
    unit = match.group(2) or "px"
    if unit == "cm":
        return Cm(amount)
    if unit == "mm":
        return Mm(amount)
    # px ベース（96dpi）
    return Inches(amount / 96)


def add_image(document: Document, base_dir: Path, src: str, width_token: str | None = None) -> None:
    src = src.strip()
    if src.startswith("./"):
        src_path = (base_dir / src[2:]).resolve()
    else:
        src_path = (base_dir / src).resolve() if not Path(src).is_absolute() else Path(src)

    if not src_path.exists():
        document.add_paragraph(f"[画像が見つかりません: {src}]")
        return

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    width = to_width(width_token)
    kwargs = {"width": width} if width else {}
    run.add_picture(str(src_path), **kwargs)

@dataclass
class MathRegion:
    page: int
    box: tuple[int, int, int, int]  # left, top, right, bottom
    score: float
    text: str


JSON_PAGE_PATTERN = re.compile(r"(?:^|_)page_(\d{3})(?:_|$)")
URL_PATTERN = re.compile(r"https?://", re.IGNORECASE)
BASE_PATTERN = re.compile(r"\([0-9]{1,3}\)\s*[0-9]{0,3}")  # (10), (12) 等
SUB_SUP_PATTERN = re.compile(r"[_^][0-9]+")
MATH_KEYWORDS = ("比率", "割合", "分数", "率", "比")


def _extract_page_number(name: str) -> int | None:
    match = JSON_PAGE_PATTERN.search(name)
    if match:
        return int(match.group(1))
    match = re.search(r"page(?:_images)?_page_(\d{3})", name)
    if match:
        return int(match.group(1))
    match = re.search(r"page_(\d{3})", name)
    if match:
        return int(match.group(1))
    return None


def _math_features(text: str) -> tuple[int, int, float, bool]:
    """Return (ops, digits, digit_ratio, has_base_marker)."""
    ops = sum(text.count(ch) for ch in "+-×÷=/%^·")
    digits = sum(ch.isdigit() for ch in text)
    length = max(1, len(text))
    digit_ratio = digits / length
    has_base = bool(BASE_PATTERN.search(text) or SUB_SUP_PATTERN.search(text))
    return ops, digits, digit_ratio, has_base


def _looks_math(text: str) -> bool:
    if URL_PATTERN.search(text):
        return False
    if any(kw in text for kw in MATH_KEYWORDS):
        return True
    ops, digits, digit_ratio, has_base = _math_features(text)
    if has_base and digits >= 4:
        return True
    if ops >= 1 and digits >= 2:
        return True
    if digit_ratio >= 0.4 and digits >= 6:
        return True
    return False


def _load_regions(
    json_path: Path,
    *,
    page: int,
    min_score: float,
    max_chars: int,
    max_aspect: float,
) -> list[MathRegion]:
    data = json.loads(json_path.read_text(encoding="utf-8"))
    regions: list[MathRegion] = []

    def add_region(box: list[int] | tuple[int, int, int, int] | None, score: float, text: str) -> None:
        if not box or len(box) != 4:
            return
        left, top, right, bottom = (int(v) for v in box)
        if right < left:
            left, right = right, left
        if bottom < top:
            top, bottom = bottom, top
        width = max(1, right - left)
        height = max(1, bottom - top)
        aspect = max(width / height, height / width)
        if aspect > max_aspect:
            return
        if len(text) > max_chars:
            return
        ops, digits, digit_ratio, has_base = _math_features(text)
        if ops < 1 and not (has_base or digit_ratio >= 0.4):
            return
        if score < min_score:
            return
        regions.append(MathRegion(page=page, box=(left, top, right, bottom), score=score, text=text))

    for para in data.get("paragraphs", []):
        text = (para.get("contents") or "").strip()
        if not text or not _looks_math(text):
            continue
        add_region(para.get("box"), float(para.get("score", 0.5)), text)

    for det in data.get("detections", []):
        text = (det.get("content") or "").strip()
        if not text or not _looks_math(text):
            continue
        points = det.get("points")
        if points and isinstance(points, list) and len(points) >= 4:
            xs = [p[0] for p in points]
            ys = [p[1] for p in points]
            box = [min(xs), min(ys), max(xs), max(ys)]
        else:
            box = det.get("box")
        score = float(det.get("rec_score", det.get("det_score", 0.5)))
        add_region(box, score, text)

    regions.sort(key=lambda r: (-r.score, r.box[1], r.box[0]))
    return regions


def _resolve_page_image(base_dir: Path, page: int) -> Path | None:
    candidate = base_dir / "page_images" / f"page_{page:03}.png"
    if candidate.exists():
        return candidate

    preprocessed = base_dir / "preprocessed"
    if preprocessed.exists():
        matches = sorted(preprocessed.glob(f"**/page_{page:03}.png"))
        if matches:
            return matches[0]

    converted = base_dir / "converted" / f"page_{page:03}.png"
    if converted.exists():
        return converted

    return None


def _build_formula_images(
    base_dir: Path,
    *,
    min_score: float,
    max_per_page: int,
    padding: int,
    max_chars: int,
    max_aspect: float,
) -> dict[int, list[Path]]:
    json_dir = base_dir / "yomi_formats" / "json"
    if not json_dir.exists():
        return {}
    json_files = sorted(json_dir.glob("*.json"))
    if not json_files:
        return {}

    figure_dir = base_dir / "figures"
    figure_dir.mkdir(parents=True, exist_ok=True)

    def crop_region(img_path: Path, region: MathRegion):
        from PIL import Image

        with Image.open(img_path) as img:
            left, top, right, bottom = region.box
            left = max(0, left - padding)
            top = max(0, top - padding)
            right = min(img.width, right + padding)
            bottom = min(img.height, bottom + padding)
            return img.crop((left, top, right, bottom)).copy()

    pages: dict[int, list[Path]] = {}
    for path in json_files:
        page = _extract_page_number(path.stem)
        if page is None and len(json_files) == 1:
            page = 1
        if page is None:
            continue
        page_image = _resolve_page_image(base_dir, page)
        if page_image is None:
            continue
        regions = _load_regions(
            path,
            page=page,
            min_score=min_score,
            max_chars=max_chars,
            max_aspect=max_aspect,
        )
        if not regions:
            continue
        picked = regions[:max_per_page]
        for idx, region in enumerate(picked, start=1):
            out_path = figure_dir / f"eq_page{page:03}_{idx:02}.png"
            cropped = crop_region(page_image, region)
            cropped.save(out_path)
            pages.setdefault(page, []).append(out_path)
    return pages


def convert_markdown(
    document: Document,
    lines: list[str],
    base_dir: Path,
    *,
    math_mode: str = "text",
) -> None:
    formula_images: dict[int, list[Path]] = {}
    if math_mode == "image":
        try:
            formula_images = _build_formula_images(
                base_dir,
                min_score=0.6,
                max_per_page=12,
                padding=6,
                max_chars=120,
                max_aspect=6.0,
            )
        except Exception:
            formula_images = {}

    paragraph_buffer: list[str] = []
    i = 0

    while i < len(lines):
        raw_line = lines[i].rstrip("\n")
        normalized = raw_line.strip()

        if not normalized:
            flush_paragraph(document, paragraph_buffer, base_dir)
            i += 1
            continue

        page_heading = PAGE_HEADING_PATTERN.match(normalized)
        if page_heading:
            flush_paragraph(document, paragraph_buffer, base_dir)
            page = int(page_heading.group("page"))
            document.add_heading(f"Page {page}", level=1)
            if math_mode == "image":
                for img_path in formula_images.get(page, []):
                    add_image(document, base_dir, str(img_path.relative_to(base_dir)))
            i += 1
            continue
        if normalized.startswith("#"):
            flush_paragraph(document, paragraph_buffer, base_dir)
            level = len(normalized) - len(normalized.lstrip("#"))
            heading_text = normalized[level:].strip()
            document.add_heading(heading_text or " ", level=min(level, 4))
            i += 1
            continue

        if is_table_line(raw_line):
            flush_paragraph(document, paragraph_buffer, base_dir)
            table_block = []
            while i < len(lines) and is_table_line(lines[i]):
                table_block.append(lines[i])
                i += 1
            add_table(document, table_block)
            continue

        if normalized.startswith("- ") or normalized.startswith("* "):
            flush_paragraph(document, paragraph_buffer, base_dir)
            paragraph = document.add_paragraph(style="List Bullet")
            render_text_content(paragraph, normalized[2:].strip())
            i += 1
            continue

        ordered_match = re.match(r"(\d+)[\.\)]\s+(.*)", normalized)
        if ordered_match:
            flush_paragraph(document, paragraph_buffer, base_dir)
            number_text = ordered_match.group(1)
            body_text = ordered_match.group(2).strip()
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{number_text}. ")
            render_text_content(paragraph, body_text)
            i += 1
            continue

        if "<img" in normalized or "![" in normalized:
            flush_paragraph(document, paragraph_buffer, base_dir)
            handled = False
            for match in IMG_HTML_PATTERN.finditer(normalized):
                add_image(document, base_dir, match.group(1), match.group(0))
                handled = True
            stripped_no_html = IMG_HTML_PATTERN.sub("", normalized)
            for match in IMG_MD_PATTERN.finditer(stripped_no_html):
                add_image(document, base_dir, match.group(1))
                handled = True
            remainder = IMG_MD_PATTERN.sub("", stripped_no_html).strip()
            if remainder:
                paragraph_buffer.append(remainder)
            if handled:
                i += 1
                continue

        paragraph_buffer.append(normalized)
        i += 1

    flush_paragraph(document, paragraph_buffer, base_dir)


def convert_file(md_path: Path, *, math_mode: str = "text") -> Path:
    if not md_path.exists():
        raise FileNotFoundError(f"Markdown ファイルが見つかりません: {md_path}")

    docx_path = md_path.with_suffix(".docx")
    document = Document()
    lines = read_markdown(md_path)
    convert_markdown(document, lines, base_dir=md_path.parent, math_mode=math_mode)
    document.save(docx_path)
    return docx_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Markdown を docx に変換する")
    parser.add_argument("markdown", nargs="?", default="merged.md", help="入力 Markdown ファイル")
    parser.add_argument(
        "--math",
        choices=["text", "image"],
        default="text",
        help="数式の扱い。text=本文としてそのまま出力（既定）、image=検出した数式領域を画像で貼る",
    )
    args = parser.parse_args()
    md_path = Path(args.markdown)
    try:
        docx_path = convert_file(md_path, math_mode=args.math)
        print(f"Word ファイルを出力しました: {docx_path}")
    except Exception as e:
        print(f"エラー: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
