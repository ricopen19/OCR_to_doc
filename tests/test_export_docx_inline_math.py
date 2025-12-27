import unittest
from pathlib import Path

from docx import Document

from export_docx import convert_markdown


class ExportDocxMathTextModeTests(unittest.TestCase):
    def test_inline_dollars_are_stripped_as_plain_text(self) -> None:
        doc = Document()
        lines = ["A $x+y$ B"]
        convert_markdown(doc, lines, base_dir=Path("."), math_mode="text")
        full_xml = "\n".join(p._p.xml for p in doc.paragraphs)
        self.assertIn("A x+y B", full_xml)
        self.assertNotIn("m:oMath", full_xml)

    def test_block_dollars_are_stripped_as_plain_text(self) -> None:
        doc = Document()
        lines = ["$$", "x+y", "$$"]
        convert_markdown(doc, lines, base_dir=Path("."), math_mode="text")
        full_xml = "\n".join(p._p.xml for p in doc.paragraphs)
        self.assertIn("x+y", full_xml)
        self.assertNotIn("$$", full_xml)


if __name__ == "__main__":
    unittest.main()
